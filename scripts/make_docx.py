from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Title
title = doc.add_heading('SmartSPA AI & Jango Virtual Try-On', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# SmartSPA Section
doc.add_heading('SmartSPA AI', level=1)

doc.add_heading('System Components', level=2)

components = [
    ('SmartSPA Agent', 'The core chatbot. It\'s a Shopify plugin that handles bookings end-to-end, from "what services do you have?" to confirming a slot.'),
    ('Service Selection Module', 'Where customers choose what they want to book. This is where every booking starts.'),
    ('Package Recommendation Engine', 'After a service is picked, up to 3 related packages appear in a carousel labeled "You might also like..." These update when users change their service selection.'),
    ('Timezone Manager', 'Everything runs on Cyprus time (Europe/Nicosia, EET). Business hours are applied to every booking automatically.'),
    ('Phone Validator', 'Expects a 9-digit Cyprus number. The system formats it to +357 without requiring the user to enter the country code.'),
    ('Room Calculator', 'Bookings are calculated based on a 2-room setup. That\'s the constraint everything works within.'),
    ('Booking Validation Engine', 'Blocks holiday bookings, rejects anything before 10am or after 8pm, and catches overlapping appointments before they cause problems.'),
    ('Conversation State Manager', 'Remembers what\'s happening in the conversation. No more losing context mid-booking.'),
]

for name, desc in components:
    doc.add_paragraph(name, style='List Bullet')
    p = doc.add_paragraph(desc)
    p.paragraph_format.space_after = Pt(6)

doc.add_heading('How It Works', level=2)

flow = [
    ('Core Flow', 'A customer lands on a Shopify store with SmartSPA installed, opens the chat, and the agent handles the whole booking from start to finish.'),
    ('Selecting a Service', 'User picks what they want: massage, treatment, whatever the business offers.'),
    ('Package Suggestions', 'Right after selecting, up to 3 related packages appear in a carousel. These update if the user changes their service choice.'),
    ('Time and Date', 'Cyprus timezone. Monday through Saturday, 10am to 8pm. Sundays are closed. The system checks for holidays before confirming anything.'),
    ('Phone Number', 'User provides a Cyprus phone number. Expects 9 digits. The system normalizes it to +357 behind the scenes.'),
    ('Room Constraint', 'The booking logic only accounts for 2 rooms. It won\'t let users book beyond that capacity.'),
    ('Conflict Detection', 'Double-booking attempts get caught. So do requests outside business hours or on holidays. The user gets told why their booking failed and can try again.'),
    ('State Management', 'The bot knows what service was chosen, what time was requested, what phone number was given. It doesn\'t lose the thread or ask for the same information twice.'),
]

for name, desc in flow:
    doc.add_paragraph(name, style='List Bullet')
    p = doc.add_paragraph(desc)
    p.paragraph_format.space_after = Pt(6)

doc.add_heading('Customer Actions', level=2)

actions = [
    'Start a Conversation — Open the chat on the Shopify store and ask anything about services, pricing, or availability.',
    'Choose a Service — Pick what they want to book from the available options.',
    'Browse Packages — Check out the suggested packages that appear after selecting a service.',
    'Provide Contact Details — Enter their phone number in 9-digit format. The system handles the rest.',
    'Select Date and Time — Pick a slot within business hours. The system validates it against holidays and operating hours.',
    'Confirm Booking — If everything checks out, the booking goes through. If not, the system explains what went wrong.',
    'Ask Follow-Up Questions — At any point, customers can ask about other services, change their booking, or request something different.',
]

for action in actions:
    doc.add_paragraph(action)

# Jango Section
doc.add_heading('Jango Virtual Try-On and Recommendation', level=1)

doc.add_heading('System Components', level=2)

jango_components = [
    ('Try It On Button', 'Embedded directly on product pages. Shows up next to shirts, pants, jackets, whatever apparel the store sells. That\'s the entry point.'),
    ('Photo Upload Module', 'Handles photo uploads from users. Specifically looks for full-body shots taken facing forward. That\'s what gives the best result.'),
    ('Virtual Try-On Engine', 'The AI brain doing the heavy lifting. Takes the user\'s photo and overlays the garment. Usually finishes in 20-25 seconds.'),
    ('Preview Viewer', 'Once processing is done, this shows the result. The user sees exactly how the item looks on them.'),
    ('Recommendation Engine', 'While the user is trying things on, this kicks in. Suggests products from different categories. Updates as they switch between items. Always shows at least 5 alternatives.'),
]

for name, desc in jango_components:
    doc.add_paragraph(name, style='List Bullet')
    p = doc.add_paragraph(desc)
    p.paragraph_format.space_after = Pt(6)

doc.add_heading('How It Works', level=2)

jango_flow = '''Customer browses the store, finds something they like, and clicks "Try It On" on the product page. The system asks for a full-body, forward-facing photo, telling the user this is what produces the best results. If they've uploaded one before, it might reuse that. The engine gets to work—takes about 20-25 seconds. If the photo isn't suitable, the system notices and asks for a clearer shot instead of trying to work with something that won't work.

After processing, the user sees a preview. They can check how the garment looks on them and decide if they want to buy. Once a product is selected, the recommendation engine starts suggesting complementary items from other categories. If the user switches to a different product, the recommendations update to match. Each customer sees at least 5 different products.'''

doc.add_paragraph(jango_flow)

doc.add_heading('Customer Actions', level=2)

jango_actions = [
    'Browse and Select — Scroll through products on the Jango Fashion store and pick something to try on.',
    'Upload a Photo — Submit a full-body photo facing forward. Or use one they\'ve uploaded before.',
    'View the Preview — See exactly how the item looks on them. Check the fit, the style, how it suits them.',
    'Make a Decision — Decide to buy based on the preview, or keep browsing if it\'s not quite right.',
    'Check Recommendations — Browse the suggested items that appear. These come from other categories and update as different products are selected.',
    'Switch Products — Try on different items to see how each looks. The recommendations refresh each time.',
]

for action in jango_actions:
    doc.add_paragraph(action)

doc.save('SmartSPA_Jango_Documentation.docx')
print('Created: SmartSPA_Jango_Documentation.docx')
